#!/usr/bin/env python3
"""Fetch full texts of court decisions referenced by EDRSR Parquet parts.

Documents are fetched politely and converted to plain text when a supported
format is encountered (RTF, HTML, PDF, DOCX, or UTF-8 text).  Conversion
libraries are optional so the metadata pipeline remains lightweight.
"""
from __future__ import annotations

import argparse
import json
import re
import time
import urllib.request
from datetime import datetime, timezone
from pathlib import Path

USER_AGENT = "JoTalbot/ukraine-edrsr-texts (+polite rate-limited mirror)"
META_COLUMNS = ("document_id", "source_url", "case_number", "court", "category", "judge", "decision_date")
SUPPORTED_EXTENSIONS = (".rtf", ".doc", ".docx", ".pdf", ".txt", ".html", ".htm")


def load_state(path: Path) -> dict:
    if path.is_file():
        state = json.loads(path.read_text(encoding="utf-8"))
        state.setdefault("fetched", [])
        state.setdefault("failed", [])
        return state
    return {"fetched": [], "failed": []}


def save_state(path: Path, state: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    state["failed"] = state.get("failed", [])[-5000:]
    path.write_text(json.dumps(state, ensure_ascii=False), encoding="utf-8")


def iter_pending(parts: list[Path], done: set[str]):
    import pyarrow.parquet as pq
    seen: set[str] = set()
    for part in parts:
        pf = pq.ParquetFile(part)
        columns = [c for c in META_COLUMNS if c in pf.schema_arrow.names]
        for batch in pf.iter_batches(batch_size=5000, columns=columns):
            for row in batch.to_pylist():
                url = (row.get("source_url") or "").strip()
                doc_id = str(row.get("document_id") or "").strip()
                if not url or not doc_id or doc_id in done or doc_id in seen:
                    continue
                if not url.lower().split("?", 1)[0].endswith(SUPPORTED_EXTENSIONS):
                    continue
                seen.add(doc_id)
                yield doc_id, url, row


def fetch(url: str, timeout: int = 60) -> bytes:
    req = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
    with urllib.request.urlopen(req, timeout=timeout) as response:
        return response.read()


def rtf_to_plain(data: bytes) -> str:
    from striprtf.striprtf import rtf_to_text
    return rtf_to_text(data.decode("utf-8", errors="replace"), errors="ignore")


def _html_to_plain(data: bytes) -> str:
    import html
    raw = data.decode("utf-8", errors="replace")
    raw = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", raw, flags=re.S | re.I)
    text = re.sub(r"<[^>]+>", " ", raw)
    return html.unescape(re.sub(r"\s+", " ", text)).strip()


def _pdf_to_plain(data: bytes) -> str:
    from io import BytesIO
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _docx_to_plain(data: bytes) -> str:
    from io import BytesIO
    from docx import Document
    document = Document(BytesIO(data))
    chunks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks).strip()


def document_text(url: str, data: bytes) -> str:
    lower = url.lower().split("?", 1)[0]
    if lower.endswith(".rtf"):
        return rtf_to_plain(data)
    if lower.endswith((".html", ".htm")):
        return _html_to_plain(data)
    if lower.endswith(".pdf"):
        return _pdf_to_plain(data)
    if lower.endswith(".docx"):
        return _docx_to_plain(data)
    if lower.endswith(".txt"):
        return data.decode("utf-8", errors="replace")
    # Legacy .doc is binary and needs an external converter (antiword/libreoffice).
    # Do not silently treat binary bytes as legal text.
    if lower.endswith(".doc"):
        raise ValueError("legacy .doc requires external text converter")
    return ""


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--parts", required=True, help="glob of EDRSR parquet parts")
    ap.add_argument("--state", required=True, help="state.json with fetched doc ids")
    ap.add_argument("--output-dir", required=True)
    ap.add_argument("--limit", type=int, default=3000)
    ap.add_argument("--delay", type=float, default=0.5)
    ap.add_argument("--max-chars", type=int, default=40000)
    args = ap.parse_args()

    import glob as globlib
    parts = sorted(globlib.glob(args.parts))
    if not parts:
        raise SystemExit(f"no parquet parts matched: {args.parts}")
    state_path = Path(args.state)
    state = load_state(state_path)
    done = set(state["fetched"])
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    rows: list[dict] = []
    fetched = failed = 0
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S")
    for doc_id, url, meta in iter_pending(parts, done):
        if fetched >= args.limit:
            break
        try:
            data = fetch(url)
            text = document_text(url, data)[: args.max_chars]
            if not text.strip():
                raise ValueError("empty text")
            rows.append({
                "document_id": doc_id,
                "case_number": meta.get("case_number"),
                "court": meta.get("court"),
                "category": meta.get("category"),
                "judge": meta.get("judge"),
                "decision_date": str(meta.get("decision_date") or "")[:10] or None,
                "source_url": url,
                "text_chars": len(text),
                "fetched_at": datetime.now(timezone.utc).isoformat(),
                "text": text,
            })
            state["fetched"].append(doc_id)
            fetched += 1
        except Exception as exc:
            state["failed"].append({"doc_id": doc_id, "url": url, "reason": str(exc)[:200]})
            failed += 1
        time.sleep(args.delay)

    shard = out_dir / f"texts-{stamp}.parquet"
    if rows:
        import pyarrow as pa
        import pyarrow.parquet as pq
        pq.write_table(pa.Table.from_pylist(rows), shard, compression="zstd")

    save_state(state_path, state)
    print(json.dumps({"fetched": fetched, "failed": failed,
                      "shard": shard.name if rows else None, "rows": len(rows),
                      "total_fetched": len(state["fetched"])}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
